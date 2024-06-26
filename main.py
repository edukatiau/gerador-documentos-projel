import os
from docx import Document
from docx2pdf import convert

from tkinter import (
    END,
    BooleanVar,
    Checkbutton,
    Tk,
    Frame,
    Label,
    Entry,
    Button,
    LEFT,
    ttk,
)
from PIL import Image, ImageTk
import tkinter.messagebox

garantia_template_file_path = "template/template_garantia.docx"  # Caminho do arquivo de template do certificado de garantia
balanceamento_template_file_path = "template/template_balanceamento.docx"  # Caminho do arquivo de template do certificado de balanceamento
vibracao_template_file_path = "template/template_vibracao.docx"  # Caminho do arquivo de template do certificado de vibração

output_garantia_path = "output/GARANTIA"  # Caminho da pasta de saída para o certificado de garantia
output_balanceamento_path = "output/BALANCEAMENTO"  # Caminho da pasta de saída para o certificado de balanceamento
output_vibracao_path = "output/VIBRAÇÃO"  # Caminho da pasta de saída para o certificado de vibração

os.makedirs(
    output_garantia_path, exist_ok=True
)  # Cria a pasta de saída, caso não exista

# Variáveis que serão substituídas no documento
dataClient = {"{CLIENTE}": "", "{OS}": "", "{NF}": ""}


# Função para substituir o texto no parágrafo
def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


class Application:
    def __init__(self, master=None):
        self.fonte = ("Verdana", "12")

        # CABEÇALHO
        self.container1 = Frame(master)
        self.container1["pady"] = 10
        self.container1["padx"] = 20
        self.container1.pack()

        image = Image.open("resource/Logo Projelmec.png")
        photo = ImageTk.PhotoImage(image)
        self.imagem = Label(self.container1, image=photo)
        self.imagem.image = photo
        self.imagem.pack()

        # OS
        self.container2 = Frame(master)
        self.container2["pady"] = 10
        self.container2["padx"] = 20
        self.container2.pack()

        self.osLabel = Label(self.container2, text="OS: ")
        self.osLabel["font"] = self.fonte
        self.osLabel.pack(side=LEFT)

        self.os = Entry(self.container2)
        self.os["width"] = 30
        self.os.pack(side=LEFT)

        # NF
        self.container3 = Frame(master)
        self.container3["pady"] = 10
        self.container3["padx"] = 20
        self.container3.pack()

        self.nfLabel = Label(self.container3, text="NF: ")
        self.nfLabel["font"] = self.fonte
        self.nfLabel.pack(side=LEFT)

        self.nf = Entry(self.container3)
        self.nf["width"] = 30
        self.nf.pack(side=LEFT)

        # CLIENTE
        self.container4 = Frame(master)
        self.container4["pady"] = 10
        self.container4["padx"] = 20
        self.container4.pack()

        self.clienteLabel = Label(self.container4, text="Cliente: ")
        self.clienteLabel["font"] = self.fonte
        self.clienteLabel.pack(side=LEFT)

        self.cliente = Entry(self.container4)
        self.cliente["width"] = 30
        self.cliente.pack(side=LEFT)

        self.teste = Frame(master)
        self.teste["pady"] = 10
        self.teste.pack()

        self.itemLabel = Label(self.teste, text="Item: ")
        self.itemLabel["font"] = self.fonte
        self.itemLabel.pack(side=LEFT)

        self.item = Entry(self.teste)
        self.item["width"] = 30
        self.item.pack(side=LEFT)

        self.quantidadeLabel = Label(self.teste, text="Quantidade: ")
        self.quantidadeLabel["font"] = self.fonte
        self.quantidadeLabel.pack(side=LEFT, padx=10)

        self.quantidade = Entry(self.teste)
        self.quantidade["width"] = 10
        self.quantidade.pack(
            side=LEFT,
        )

        self.varVibracao = BooleanVar()
        c1 = Checkbutton(
            self.teste,
            text="Vibração",
            variable=self.varVibracao,
            onvalue=True,
            offvalue=False,
        )
        c1.pack(side=LEFT, padx=10)

        self.addBtn = Button(self.teste)
        self.addBtn["text"] = "Adicionar"
        self.addBtn["font"] = ("Calibri", "12")
        self.addBtn["width"] = 10
        self.addBtn["command"] = self.add_item
        self.addBtn.pack(side=LEFT, padx=10)
        
        self.deleteBtn = Button(self.teste)
        self.deleteBtn["text"] = "Deletar itens"
        self.deleteBtn["font"] = ("Calibri", "12")
        self.deleteBtn["width"] = 10
        self.deleteBtn["command"] = self.delete_item
        self.deleteBtn.pack(side=LEFT, padx=10)

        self.i = 1

        self.treeview = ttk.Treeview(columns=("EQUIPAMENTO", "QUANTIDADE", "VIBRAÇÃO"))
        self.treeview.heading("#0", text="ITEM")
        self.treeview.heading("EQUIPAMENTO", text="EQUIPAMENTO")
        self.treeview.heading("QUANTIDADE", text="QUANTIDADE")
        self.treeview.heading("VIBRAÇÃO", text="VIBRAÇÃO")
        self.treeview.pack()

        # BOTÃO
        self.container5 = Frame(master)
        self.container5["pady"] = 20
        self.container5["padx"] = 20
        self.container5.pack()

        self.gerar = Button(self.container5)
        self.gerar["text"] = "Gerar Documentos"
        self.gerar["font"] = ("Calibri", "22")
        self.gerar["width"] = 20
        self.gerar["command"] = self.gerar_doc
        self.gerar.pack()

    # Função para adicionar item na treeview
    def add_item(self):
        if self.item.get() == "" or self.quantidade.get() == "":
            tkinter.messagebox.showerror("Erro", "Preencha todos os campos!")
            return
        elif not self.quantidade.get().isdigit():
            tkinter.messagebox.showerror(
                "Erro", "O campo quantidade deve ser numérico!"
            )
            return
        else:
            item = self.item.get()
            quantidade = self.quantidade.get()
            vibracao = self.varVibracao.get()
            self.treeview.insert(
                "", END, text=str(self.i), values=(item, quantidade, vibracao)
            )
            self.i += 1

    # Função para deletar item na treeview
    def delete_item(self):
        for item in self.treeview.get_children():
            self.treeview.delete(item)
            self.i = 1

    # Função para receber os dados do cliente
    def recebe_dataClient(self):
        if self.cliente.get() == "" or self.os.get() == "" or self.nf.get() == "":
            tkinter.messagebox.showerror("Erro", "Preencha todos os campos!")
            return False
        elif not self.os.get().isdigit() or not self.nf.get().isdigit():
            tkinter.messagebox.showerror(
                "Erro", "Os campos OS e NF devem ser numéricos!"
            )
            return False
        elif len(self.os.get()) < 5:
            tkinter.messagebox.showerror(
                "Erro", "O campo OS deve ter pelo menos 5 dígitos!"
            )
            return False
        elif len(self.nf.get()) < 5:
            tkinter.messagebox.showerror(
                "Erro", "O campo NF deve ter pelo menos 5 dígitos!"
            )
            return False
        elif len(self.cliente.get()) < 5:
            tkinter.messagebox.showerror(
                "Erro", "O campo Cliente deve ter pelo menos 5 caracteres!"
            )
            return False
        elif len(self.cliente.get()) > 50:
            tkinter.messagebox.showerror(
                "Erro", "O campo Cliente deve ter no máximo 50 caracteres!"
            )
            return False
        elif len(self.os.get()) > 6:
            tkinter.messagebox.showerror(
                "Erro", "O campo OS deve ter no máximo 6 caracteres!"
            )
            return False
        elif len(self.nf.get()) > 6:
            tkinter.messagebox.showerror(
                "Erro", "O campo NF deve ter no máximo 6 caracteres!"
            )
            return False

        nf = self.nf.get()
        nf = nf[0:2] + "." + nf[2:6]

        dataClient["{CLIENTE}"] = self.cliente.get()
        dataClient["{OS}"] = self.os.get()
        dataClient["{NF}"] = nf
        return True

    # Função para gerar o documento
    def gerar_doc(self):
        # Recebe os dados do cliente
        if not self.recebe_dataClient():
            return

        # Gera o documento de garantia
        self.doc_garantia()

        # Recebe os dados do balanceamento
        if not self.recebe_dataPedido():
            return

        # Gera o documento de balanceamento
        self.doc_balanceamento()

        tkinter.messagebox.showinfo("Sucesso", "Documento gerado com sucesso!")

    # Função para gerar o documento de garantia
    def doc_garantia(self):
        garantia_template = Document(garantia_template_file_path)

        # Substitui as variáveis no documento
        for variable_key, variable_value in dataClient.items():
            for paragraph in garantia_template.paragraphs:
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in garantia_template.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(
                                paragraph, variable_key, variable_value
                            )

        try:
            # Salva o documento
            output_garantia_file_path = "{}/CG-OS-{}.docx".format(
                output_garantia_path, dataClient["{OS}"]
            )
            
            if (os.path.exists(output_garantia_file_path)):
                output_garantia_file_path = "{}/CG-OS-{} NF {}.docx".format(
                    output_garantia_path, dataClient["{OS}"], dataClient["{NF}"]
                )
            
            garantia_template.save(output_garantia_file_path)
            print("Docx Done!")

            # Converte o documento para PDF
            os.makedirs(output_garantia_path + "/PDF", exist_ok=True)
            
            if os.path.exists(output_garantia_file_path):
                pdf_file_path = "{}/PDF/CG-OS-{} NF {}.pdf".format(output_garantia_path, dataClient["{OS}"], dataClient["{NF}"])
            else:
                pdf_file_path = "{}/PDF/CG-OS-{}.pdf".format(output_garantia_path, dataClient["{OS}"])

            # Realiza a conversão para PDF usando o caminho determinado
            convert(output_garantia_file_path, pdf_file_path)
                
            print("PDF Done!")

            tkinter.messagebox.showinfo("Sucesso", "Documento gerado com sucesso!")

        except Exception as e:
            print("Error: ", e)
            tkinter.messagebox.showerror("Erro", "Erro ao gerar o documento!")

    # Função para receber os dados do balanceamento
    def recebe_dataPedido(self):
        self.itens = []

        for child in self.treeview.get_children():
            item = self.treeview.item(child)["text"]
            equipamento = self.treeview.item(child)["values"][0]
            quantidade = self.treeview.item(child)["values"][1]
            hasVibracao = self.treeview.item(child)["values"][2]

            self.itens.append(
                {
                    "{ITEM}": item,
                    "{EQUIPAMENTO}": equipamento,
                    "{QUANTIDADE}": quantidade,
                    "{VIBRAÇÃO}": hasVibracao,
                }
            )

        return True

    def doc_balanceamento(self):
        try:
            # Recebe os dados do cliente
            if not self.recebe_dataClient():
                return

            i = 1  # contador para os itens
            for item in self.itens:
                for chave, valor in item.items():
                    if chave == "{QUANTIDADE}":
                        quantidade = int(valor)
                        # Para cada quantidade do item, criar um documento
                        for j in range(quantidade):
                            balanceamento_template = Document(
                                balanceamento_template_file_path
                            )
                            for paragraph in balanceamento_template.paragraphs:
                                replace_text_in_paragraph(paragraph, "ITEM", str(i))
                                replace_text_in_paragraph(paragraph, "SUB", str(j + 1))
                                replace_text_in_paragraph(
                                    paragraph, "{EQUIPAMENTO}", item["{EQUIPAMENTO}"]
                                )
                                for variable_key, variable_value in dataClient.items():
                                    replace_text_in_paragraph(
                                        paragraph, variable_key, variable_value
                                    )

                            try:
                                # Salva o documento
                                os.makedirs(output_balanceamento_path, exist_ok=True)
                                output_balanceamento_file_path = (
                                    "{}/CB-OS-{}-{}.docx".format(
                                        output_balanceamento_path,
                                        dataClient["{OS}"],
                                        str(i) + "." + str(j + 1),
                                    )
                                )
                                balanceamento_template.save(
                                    output_balanceamento_file_path
                                )
                                print("Docx Done!")

                                # Converte o documento para PDF
                                os.makedirs(
                                    output_balanceamento_path + "/PDF", exist_ok=True
                                )
                                convert(
                                    output_balanceamento_file_path,
                                    "{}/PDF/CB-OS-{}-{}.pdf".format(
                                        output_balanceamento_path,
                                        dataClient["{OS}"],
                                        str(i) + "." + str(j + 1),
                                    ),
                                )
                                print("PDF Done!")

                            except Exception as e:
                                print("Error: ", e)
                                tkinter.messagebox.showerror(
                                    "Erro", "Erro ao gerar o documento!"
                                )

                    if chave == "{VIBRAÇÃO}" and valor == "True":
                        self.doc_vibracao(item)

                i += 1

        except Exception as e:
            print("Error: ", e)
            tkinter.messagebox.showerror("Erro", "Erro ao gerar o documento!")

    def doc_vibracao(self, item):
        
        i = 1
        for chave, valor in item.items():
            if chave == "{QUANTIDADE}":
                quantidade = int(valor)
                # Para cada quantidade do item, criar um documento
                for j in range(quantidade):
                                        
                    self.janela_vibracao(item, i, j)
                                        
                    vibracoes = {"{VIB1}": self.vib1, "{VIB2}": self.vib2}
                    
                    vibracao_template = Document(vibracao_template_file_path)
                    for paragraph in vibracao_template.paragraphs:
                        replace_text_in_paragraph(paragraph, "ITEM", str(i))
                        replace_text_in_paragraph(paragraph, "SUB", str(j + 1))
                        replace_text_in_paragraph(paragraph, "{EQUIPAMENTO}", item["{EQUIPAMENTO}"])
                        
                        for variable_key, variable_value in dataClient.items():
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)
                        
                        for variable_key, variable_value in vibracoes.items():
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)

                    try:
                        # Salva o documento
                        os.makedirs(output_vibracao_path, exist_ok=True)
                        output_vibracao_file_path = "{}/CV-OS-{}-{}.docx".format(
                            output_vibracao_path,
                            dataClient["{OS}"],
                            str(i) + "." + str(j + 1),
                        )
                        vibracao_template.save(output_vibracao_file_path)
                        print("Docx Done!")

                        # Converte o documento para PDF
                        os.makedirs(output_vibracao_path + "/PDF", exist_ok=True)
                        convert(
                            output_vibracao_file_path,
                            "{}/PDF/CV-OS-{}-{}.pdf".format(
                                output_vibracao_path,
                                dataClient["{OS}"],
                                str(i) + "." + str(j + 1),
                            ),
                        )
                        print("PDF Done!")

                    except Exception as e:
                        print("Error: ", e)
                        tkinter.messagebox.showerror(
                            "Erro", "Erro ao gerar o documento!"
                        )

        pass

    def janela_vibracao(self, item, i, j):
        self.janela = Tk()
        self.janela.title("Vibrações")
        self.janela.geometry("300x200")
        
        self.itemLabel = Label(self.janela, text="Item: " + str(i) + "." + str(j + 1))
        self.itemLabel["font"] = self.fonte
        self.itemLabel.pack()
        self.eqpLabel = Label(self.janela, text="Equipamento: " + item["{EQUIPAMENTO}"])
        self.eqpLabel["font"] = self.fonte
        self.eqpLabel.pack()
        
        self.vib1Label = Label(self.janela, text="Vibração 1: ")
        self.vib1Label.pack()
        
        self.vibracao1 = Entry(self.janela)
        self.vibracao1.pack()
        
        self.vib2Label = Label(self.janela, text="Vibração 2: ")
        self.vib2Label.pack()
        
        self.vibracao2 = Entry(self.janela)
        self.vibracao2.pack()
        
        self.confirmar = Button(self.janela)
        self.confirmar["text"] = "Confirmar"
        self.confirmar["width"] = 10
        self.confirmar["command"] = self.confirmar_vibracao
        self.confirmar.pack(pady=10)
        
        self.janela.mainloop()

    def confirmar_vibracao(self):
        self.vib1 = self.vibracao1.get()
        self.vib2 = self.vibracao2.get()
        
        self.janela.quit()
        self.janela.destroy()
        
        
# Inicializa a aplicação
root = Tk()
Application(root)
root.title("Gerador de Certificados Projelmec")
root.mainloop()
